import docx 
from transformers import pipeline
from fpdf import FPDF

#imports the document
chronology1 = docx.Document('Chron1.docx')
text = ""
for para in chronology1.paragraphs:
    text += para.text

def docx_to_string (file_path):
    # load the document
    doc = docx.Document(file_path)
    
    # initialize an empty string to store the text
    full_text = ""
    
    # iterate through the paragraphs in the document
    for para in doc.paragraphs:
        full_text += para.text + "\n"  # add a newline character to separate paragraphs
    
    return full_text
# converts to string
file_path = "Chron1.docx"
document_text = docx_to_string(file_path)

#spilts it by paragraph
line = document_text.splitlines()

# load model
ner = pipeline("ner", model ="dslim/bert-base-NER")

# extract entities
entities = ner(document_text)
print("Extracted Entities:", entities)

# filter for dates and times
date_time_entities = [entity for entity in entities if entity['entity'] in ['DATE','TIME']]
print("Date and Time Entities:", date_time_entities)


# function to extract context around the identified dates and times
def extract_event_context(text, entities, window=50):
    events = []
    for entity in entities:
        start = max(0, entity['start'] - window)
        end = min(len(text), entity['end'] + window)
        event_context = text[start:end]
        events.append((entity['word'], event_context))
    return events

# extract context around the dates and times
event_contexts = extract_event_context(text, date_time_entities)
print("Event Contexts:", event_contexts)

# format the output as date, time: event
formatted_events = []
current_date = ""
current_time = ""

for word, context in event_contexts:
    if "DATE" in word:
        current_date = word
    elif "TIME" in word:
        current_time = word
        # remove the date and time from the context
        event_description = context.replace(current_date, "").replace(current_time, "").strip()
        formatted_events.append(f"{current_date}, {current_time}: {event_description}")

print("Formatted Events:", formatted_events)

#create a pdf document
pdf = FPDF()
pdf.add_page()
pdf.set_font("Arial", size=12)

#add events into the pdf
for event in formatted_events:
    pdf.multi_cell(0,10,event)

pdf.output("final_chronology.pdf")

print("PDF file created successfully.")